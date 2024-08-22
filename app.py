# streamlit_app/app.py
# (Additional import)
from utils import data_mapping, postprocessing, preprocessing
from models.segmentation_model import load_segmentation_model, segment_image
from models.text_extraction_model import extract_text_with_otsu
from models.summarization_model import summarize_text
import streamlit as st
import os
from PIL import Image
import docx
from docx import Document
import json
import pandas as pd
import matplotlib.pyplot as plt

# Setting up paths
DATA_DIR = "data"
INPUT_DIR = os.path.join(DATA_DIR, "input_images")
OUTPUT_DIR = os.path.join(DATA_DIR, "output")
SEGMENTED_DIR = os.path.join(OUTPUT_DIR, "segmented_objects")
TEXT_EXTRACTION_DIR = os.path.join(OUTPUT_DIR, "text_extraction")
SUMMARIZATION_DIR = os.path.join(OUTPUT_DIR, "summarization")
REPORT_DIR = os.path.join(OUTPUT_DIR,"report")
# Create directories if they don't exist
os.makedirs(INPUT_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(SEGMENTED_DIR, exist_ok=True)
os.makedirs(SUMMARIZATION_DIR, exist_ok=True)
os.makedirs(TEXT_EXTRACTION_DIR, exist_ok=True)
os.makedirs(REPORT_DIR, exist_ok=True)
st.title("AI Pipeline for Image Segmentation and Object Analysis")


# Add this function in streamlit_app/app.py
def upload_image():
    st.header("Upload an Image")

    uploaded_file = st.file_uploader("Choose an image...", type=["png", "jpg", "jpeg"])

    if uploaded_file is not None:
        image = Image.open(uploaded_file)
        st.image(image, caption='Uploaded Image', use_column_width=True)

        # Save the uploaded image to INPUT_DIR
        image.save(os.path.join(INPUT_DIR, uploaded_file.name))
        st.success(f"Image saved successfully as {uploaded_file.name} in {INPUT_DIR}.")


# (Adding to the run_segmentation function)
# streamlit_app/app.py

def run_segmentation_and_identification():
    st.header("Segmentation and Object Identification")

    # Check if there's an uploaded image
    image_files = os.listdir(INPUT_DIR)
    if not image_files:
        st.warning("Please upload an image first!")
        return

    # Load the uploaded image
    image_path = os.path.join(INPUT_DIR, image_files[-1])
    image = Image.open(image_path)
    st.image(image, caption="Original Image", use_column_width=True)

    if st.button("Segment and Identify"):

        image_tensor = preprocessing.preprocess_image(image)

        # Load the models
        model_seg = load_segmentation_model()

        print("Running segmentation and object identification...")

        # Run segmentation
        scores, boxes, classes = segment_image(model_seg, image_tensor)

        # Save the segmented objects and perform identification
        segmented_images = postprocessing.save_segmented_objects(image, boxes, classes, SEGMENTED_DIR)

        # Load COCO labels from the text file
        coco_labels_file_path = r"C:\Users\Aditya PC\PycharmProjects\wasserstoff_tasks\coco_labels.txt"
        with open(coco_labels_file_path, "r") as file:
            coco_labels = file.read().splitlines()

        for idx, (segmented_image_path, class_idx, score) in enumerate(zip(segmented_images, classes, scores)):
            segmented_image = Image.open(segmented_image_path)

            # Ensure class_idx is converted to an integer
            class_idx = int(class_idx)

            # Map class index to COCO label name
            label_name = coco_labels[class_idx - 1]  # Subtracting 1 because class indices are usually 1-based

            # Display the segmented image with the label name and score
            st.image(segmented_image, caption=f"Segment {idx + 1}: {label_name} ({score * 100:.2f}%)",
                     use_column_width=True)

        st.success("Segmentation and object identification completed.")

        # Save the results if needed (e.g., save identified classes to a file)


# streamlit_app/app.py
# (Additional import)
def run_text_extraction():
    st.header("Text Extraction from Segmented Images")

    # Get all segmented images from the SEGMENTED_DIR
    segmented_images = [img for img in os.listdir(SEGMENTED_DIR) if img.endswith((".png", ".jpg", ".jpeg"))]

    # Path for the text extraction doc
    text_extraction_path = os.path.join(TEXT_EXTRACTION_DIR, "text_extraction.doc")

    if not segmented_images:
        st.warning("No segmented images found. Please run segmentation first.")
        return

    # Check if text_extraction.doc exists, if not create it
    if not os.path.exists(text_extraction_path):
        doc = docx.Document()
        doc.add_heading('Text Extraction', 0)
        doc.save(text_extraction_path)

    # Initialize session state to store extracted texts if it doesn't exist
    if "extracted_texts" not in st.session_state:
        st.session_state.extracted_texts = {}

    # Iterate through each segmented image
    for image_name in segmented_images:
        # Display the image with its name
        image_path = os.path.join(SEGMENTED_DIR, image_name)
        image = Image.open(image_path)
        st.image(image, caption=image_name, use_column_width=True)

        # Add a button to extract text for this image
        if st.button(f"Extract Text for {image_name}"):
            with st.spinner(f"Extracting text for {image_name}..."):
                # Call the text extraction function
                extracted_text = extract_text_with_otsu(image_path)

                # Open the existing document
                doc = Document(text_extraction_path)

                # Add image to the doc
                doc.add_picture(image_path)
                # Add image name and extracted text to the doc
                doc.add_paragraph(f"Image Name: {image_name}")
                doc.add_paragraph(f"Extracted Text:\n{extracted_text}")

                # Save the updated doc
                doc.save(text_extraction_path)
                st.success(f"Text extracted and appended to {text_extraction_path}.")

                # Store the extracted text in session state to persist it across interactions
                st.session_state.extracted_texts[image_name] = extracted_text

    st.write("Text extraction completed.")
    return st.session_state.extracted_texts


def run_summarization():
    st.header("Image Summarization")

    # List the segmented images from the directory
    segmented_images = os.listdir(SEGMENTED_DIR)

    if not segmented_images:
        st.warning("No segmented images found! Please run segmentation first.")
        return

    # Path for the summarization doc
    doc_file_path = os.path.join(SUMMARIZATION_DIR, "image_summarization.doc")

    # Check if the document exists, create it if it doesn't
    if not os.path.exists(doc_file_path):
        doc = docx.Document()
        doc.add_heading('Image Summarization', 0)
        doc.save(doc_file_path)

    # Initialize session state to store summaries if it doesn't exist
    if "summaries" not in st.session_state:
        st.session_state.summaries = {}

    # Iterate over each segmented image
    for image_file in segmented_images:
        image_path = os.path.join(SEGMENTED_DIR, image_file)
        image = Image.open(image_path)
        st.image(image, caption=f"Segmented Image: {image_file}", use_column_width=True)

        # Add the "Run Summarization" button for each image
        if st.button(f"Run Summarization for {image_file}"):
            with st.spinner(f"Summarizing {image_file}..."):
                # Generate a description using the summarization function
                description = summarize_text(image_path)

                # Open the existing document
                doc = Document(doc_file_path)

                # Add image and description to the document
                doc.add_picture(image_path)
                doc.add_paragraph(f"Image Name: {os.path.basename(image_path)}")
                doc.add_paragraph(f"Description: {description}")

                # Save the updated document
                doc.save(doc_file_path)

                # Store the description in session state to persist it across interactions
                st.session_state.summaries[image_file] = description

                st.success(f"Summarization complete for {image_file}!")

    return st.session_state.summaries


def run_final_report(extracted_texts, summaries):
    st.header("Final Report Generation")

    # Load all segmented images from SEGMENTED_DIR
    segmented_images = sorted(os.listdir(SEGMENTED_DIR))
    # st.write(f"Debug: Loaded {len(segmented_images)} segmented images from {SEGMENTED_DIR}: {segmented_images}")
    #
    # # Debug: Display contents of extracted_texts and summaries before mapping
    # st.write(f"Debug: Extracted Texts Dictionary: {extracted_texts}")
    # st.write(f"Debug: Summaries Dictionary: {summaries}")

    # Initialize data map
    data_map = {}

    # Map data to objects using image names
    for image_name in segmented_images:
        description = summaries.get(image_name, None)
        extracted_text = extracted_texts.get(image_name, None)

        # Debugging statements for data mapping
        # st.write(f"Debug: Mapping data for image: {image_name}")
        # st.write(f"Debug: Description: {description}")
        # st.write(f"Debug: Extracted Text: {extracted_text}")

        # Store data in the map
        data_map[image_name] = {
            "description": description,
            "extracted_text": extracted_text
        }

    # Save the data structure as a JSON file for future reference
    json_path = os.path.join(REPORT_DIR, 'data_mapping.json')
    with open(json_path, 'w') as json_file:
        json.dump(data_map, json_file, indent=4)
    st.success(f"Data mapping saved to {json_path}")
    # st.write(f"Debug: Data mapping saved as JSON to {json_path}")

    # Output the final image with annotations using Matplotlib
    fig, ax = plt.subplots(1, 1, figsize=(10, 10))
    original_image_path = os.path.join(INPUT_DIR, os.listdir(INPUT_DIR)[0])
    original_image = Image.open(original_image_path)
    # st.write(f"Debug: Loaded original image from {original_image_path}")
    ax.imshow(original_image)
    ax.set_title("Original Image with Segmented Object Annotations")

    # Annotate each segment on the image
    for idx, image_name in enumerate(segmented_images):
        # Debugging annotation positions
        # st.write(f"Debug: Annotating image {image_name} at position: {10 + idx * 20}, 20")
        ax.annotate(f"Object {idx + 1}", xy=(10 + idx * 20, 20), color="red", fontsize=12)

    # Display the final annotated image
    st.pyplot(fig)
    # st.write("Debug: Displayed final annotated image")

    # Create a DataFrame to store the mapped data
    df_data_map = {
        "Object ID": [],
        "Description": [],
        "Extracted Text": [],
        "Image Name": []
    }

    # Fill in the DataFrame with data from the mapped objects
    for idx, image_name in enumerate(segmented_images):
        # st.write(f"Debug: Processing data for image: {image_name}")

        df_data_map["Object ID"].append(idx + 1)  # Start Object ID from 1
        df_data_map["Description"].append(data_map[image_name].get("description", None))
        df_data_map["Extracted Text"].append(data_map[image_name].get("extracted_text", None))
        df_data_map["Image Name"].append(image_name)

        # st.write(f"Debug: Object {idx + 1}: Description - {data_map[image_name].get('description', None)}")
        # st.write(f"Debug: Object {idx + 1}: Extracted Text - {data_map[image_name].get('extracted_text', None)}")

    # Convert to a pandas DataFrame
    df = pd.DataFrame(df_data_map)
    # st.write(f"Debug: Created DataFrame for data mapping: {df.head()}")

    # Display the table summarizing all data
    st.write("Data Summary for Segmented Objects")
    st.dataframe(df)

    # Optionally, save this table to a CSV file
    csv_path = os.path.join(REPORT_DIR, 'data_summary.csv')
    df.to_csv(csv_path, index=False)
    st.success(f"Data summary saved to {csv_path}")
    # st.write(f"Debug: Data summary saved to {csv_path}")

    st.success("Final report generation completed.")


# Make sure to include this final step in the Streamlit app layout
def main():
    st.title("Wasserstoff AI Internship Task")

    # Inside the main() function in streamlit_app/app.py

    # Initialize session state variables if not present
    if 'extracted_texts' not in st.session_state:
        st.session_state.extracted_texts = {}
    if 'summaries' not in st.session_state:
        st.session_state.summaries = {}

    # Option to select task from the sidebar
    option = st.sidebar.selectbox("Choose a task",
                                  ["Upload Image", "Segment and Identify Objects",
                                   "Text Extraction", "Image Summarization", "Final Report"])

    # Task 1: Image Upload
    if option == "Upload Image":
        upload_image()

    # Task 2: Segmentation and Identification
    elif option == "Segment and Identify Objects":
        run_segmentation_and_identification()

    # Task 3: Text Extraction
    elif option == "Text Extraction":
        st.session_state.extracted_texts = run_text_extraction()  # Store extracted texts in session state

    # Task 4: Image Summarization
    elif option == "Image Summarization":
        st.session_state.summaries = run_summarization()  # Store summaries in session state

    # Task 5: Final Report Generation
    elif option == "Final Report":
        # Check if extracted texts and summaries are already in session state
        if not st.session_state.extracted_texts:
            st.write("Run Text Extraction before generating the report...")
        elif not st.session_state.summaries:
            st.write("Run Image Summarization before generating the report...")
        else:
            # Run the final report with both extracted texts and summaries
            run_final_report(st.session_state.extracted_texts, st.session_state.summaries)


if __name__ == "__main__":
    main()
